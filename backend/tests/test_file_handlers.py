"""
Tests for the document masking handlers (PDF / DOCX / image).

These verify the exact-redaction guarantees that the old implementations
violated:

- PDFs: PII split across lines must still be redacted (the old
  ``page.search_for`` silently failed on multi-line matches and left the PII
  visible while the report claimed it was masked).
- Images: only the exact flagged words may be blacked out (the old
  implementation redacted every OCR box whose lowercase text matched any
  token of a flagged PII string — an email's "com" blacked out every
  ".com" in the image).
- DOCX: replacement must happen at run granularity so formatting survives,
  and headers/footers must be scanned (they were never checked before).

Regex-backed entities only (EMAIL_ADDRESS / PHONE_NUMBER) — the spaCy models
are lazy-loaded, so these run in CI without downloading them.
"""
import io
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import cv2
import numpy as np

from backend.app import file_handlers, pii_engine

ENTITIES = ["EMAIL_ADDRESS", "PHONE_NUMBER"]
STYLE = "LABEL"


def _make_pdf(lines, filename=None):
    import fitz
    doc = fitz.open()
    page = doc.new_page(width=595, height=300)
    y = 60
    for line in lines:
        page.insert_text((72, y), line, fontsize=12)
        y += 24
    data = doc.tobytes()
    doc.close()
    return data


def test_pdf_masks_pii_split_across_lines():
    """A phone number wrapped across two PDF lines must be fully redacted.

    Old behavior: page.search_for("+91 98765\\n43210") fails to locate the
    multi-line match, so the digits stayed visible in the output while the
    report still claimed they were masked. New behavior: word-level box
    redaction keys off the spans themselves, so both halves get blacked out.
    """
    import fitz
    pdf = _make_pdf(["Call +91 98765", "43210 today"])
    out_bytes, report = file_handlers.process_pdf(pdf, ENTITIES, pii_engine.detect_raw)

    with fitz.open(stream=out_bytes, filetype="pdf") as doc:
        text = doc[0].get_text()

    assert "98765" not in text
    assert "43210" not in text
    assert "Call" in text
    assert "today" in text
    assert any("PHONE_NUMBER" in item["pii_types"] for item in report)


def test_pdf_does_not_redact_words_containing_a_flagged_substring():
    """A word that merely CONTAINS the flagged substring must survive.

    The old search_for-based code redacted every occurrence of the flagged
    substring — flagging "John" would also black out "Johnston". With
    word-exact mapping only the flagged words' boxes are redacted.
    """
    # AADHAAR-like 12-digit number is not used here; instead prove the point
    # with an email whose local part also appears inside another word.
    pdf = _make_pdf(["Contact support@corp.com", "support teams work here"])
    out_bytes, report = file_handlers.process_pdf(pdf, ENTITIES, pii_engine.detect_raw)

    import fitz
    with fitz.open(stream=out_bytes, filetype="pdf") as doc:
        text = doc[0].get_text()

    assert "support@corp.com" not in text
    assert "support teams work here" in text
    assert any("EMAIL_ADDRESS" in item["pii_types"] for item in report)


def test_docx_preserves_formatting_and_masks_runs():
    """Run-level redaction must mask the PII but keep bold/italic formatting."""
    from docx import Document
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Call me at ")
    r2 = p.add_run("+91 98765 43210")
    r2.bold = True
    r3 = p.add_run(" today.")
    r3.italic = True

    out_bytes, report = file_handlers.process_docx(
        _docx_bytes(doc), ENTITIES, pii_engine.detect_raw, STYLE
    )

    result = Document(io.BytesIO(out_bytes))
    para = result.paragraphs[0]
    assert para.text == "Call me at [PHONE_NUMBER_MASKED] today."
    # Formatting preserved: the masked token is still bold, the tail still italic.
    runs_by_text = {r.text: r for r in para.runs}
    assert runs_by_text["[PHONE_NUMBER_MASKED]"].bold is True
    assert runs_by_text[" today."].italic is True
    assert runs_by_text["Call me at "].bold is None
    assert any("PHONE_NUMBER" in item["pii_types"] for item in report)


def test_docx_scans_headers_and_footers():
    """PII in the document header/footer must be masked (was never scanned)."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("Body text with no PII.")
    header = doc.sections[0].header
    header.paragraphs[0].add_run("Support: support@corp.com")

    out_bytes, report = file_handlers.process_docx(
        _docx_bytes(doc), ENTITIES, pii_engine.detect_raw, STYLE
    )

    result = Document(io.BytesIO(out_bytes))
    header_text = result.sections[0].header.paragraphs[0].text
    assert "support@corp.com" not in header_text
    assert "[EMAIL_ADDRESS_MASKED]" in header_text
    assert any("EMAIL_ADDRESS" in item["pii_types"] for item in report)


def test_image_masks_email_but_not_unrelated_occurrences():
    """Only the exact flagged words may be blacked out.

    Old behavior: every OCR box whose lowercase text was a token of the
    flagged PII (e.g. "com" from "john.doe@example.com") was redacted — so a
    second, legitimate "example.com" in the same image was blacked out too.
    """
    engine_ocr = _ocr_engine()
    img = np.full((220, 900, 3), 255, np.uint8)
    cv2.putText(img, "Email john.doe@example.com today.", (20, 70), cv2.FONT_HERSHEY_SIMPLEX, 0.9, (0, 0, 0), 2)
    cv2.putText(img, "Visit example.com for info.", (20, 150), cv2.FONT_HERSHEY_SIMPLEX, 0.9, (0, 0, 0), 2)
    ok, png = cv2.imencode(".png", img)
    assert ok

    masked_bytes, report = file_handlers.mask_pii_in_image(png.tobytes(), ENTITIES, pii_engine.detect_raw)

    assert any("EMAIL_ADDRESS" in item["pii_types"] for item in report)

    masked_img = cv2.imdecode(np.frombuffer(masked_bytes, np.uint8), cv2.IMREAD_COLOR)
    ocr_after, _ = engine_ocr(masked_img)
    text_after = " ".join(r[1] for r in ocr_after)
    # The flagged email must be gone from the masked image…
    assert "john.doe" not in text_after
    # …while the legitimate second occurrence survives.
    assert "example.com" in text_after


def _docx_bytes(doc):
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _ocr_engine():
    pytest = __import__("pytest")
    pytest.importorskip("rapidocr_onnxruntime")
    return file_handlers._get_ocr_engine()
