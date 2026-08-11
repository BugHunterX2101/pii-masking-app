import io
import re
from typing import Optional

import fitz  # PyMuPDF
from docx import Document

from backend.app import pii_engine


# Contact signal for the document-header name rule: an email, or a phone-like
# string of 7+ digits. A resume's first line ("Vedit Agrawal") is only treated
# as a person name when the following lines actually carry contact info —
# "Job Description" headers and cover-letter salutations never match.
_CONTACT_IN_LINE_RE = re.compile(
    r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}"
    r"|(?:\+?\d[\d\s\-().]{6,}\d)"
)


def _pdf_name_lines(raw_words, word_spans) -> list:
    """Header-block name detection for native-text PDFs.

    The first line of the document, if it reads like a person name (2-4
    title-cased words) and the next up-to-3 lines carry an email or phone, is
    returned as a PERSON candidate span. This is how a resume's name line is
    caught even though spaCy's NER misses it."""
    groups = {}
    order = []
    for wi, w in enumerate(raw_words):
        key = (w[5], w[6])  # (block, line)
        if key not in groups:
            groups[key] = []
            order.append(key)
        groups[key].append(wi)
    if not order:
        return []
    first = groups[order[0]]
    if not pii_engine.name_line_tokens_ok([raw_words[wi][4] for wi in first]):
        return []
    contact_words = []
    for key in order[1:4]:
        for wi in groups[key]:
            contact_words.append(raw_words[wi][4])
    if not _CONTACT_IN_LINE_RE.search(" ".join(contact_words)):
        return []
    return [(word_spans[first[0]][0], word_spans[first[-1]][1])]

# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

_ocr_engine = None


def _get_ocr_engine():
    """Lazy singleton: RapidOCR models load once per process (~6s cold).

    RapidOCR is a local ONNX OCR (no cloud API, no credentials, works on any
    host including Hugging Face Spaces). The GCP Vision path it replaces
    required GOOGLE_APPLICATION_CREDENTIALS and failed on deployments without
    them — and added a network round-trip per image.
    """
    global _ocr_engine
    if _ocr_engine is None:
        from rapidocr_onnxruntime import RapidOCR
        _ocr_engine = RapidOCR()
    return _ocr_engine


def _mask_token(entity_type: str, masking_style: str) -> str:
    """Replacement token, kept byte-for-byte identical to detect_and_mask_text."""
    style = (masking_style or "LABEL").upper()
    if style == "BLACKOUT":
        return "████████"
    if style == "ASTERISK":
        return "***"
    return f"[{entity_type}_MASKED]"


def _spans_to_boxes(text: str, words, spans) -> set:
    """Map detection spans (in `text` space) to word indices.

    `words` is a list of (start, end, index) tuples describing where each word
    sits inside `text`. Returns the set of word indices whose characters
    overlap any detection span — partial overlaps included, so a span that
    ends mid-word still redacts the whole word (no partial characters left
    visible), while words merely *containing* a substring are never touched
    unless they actually intersect the flagged span.
    """
    redacted = set()
    for ws, we, wi in words:
        for s, e in spans:
            if ws < e and we > s:
                redacted.add(wi)
                break
    return redacted


# ---------------------------------------------------------------------------
# PDF (native text) — word-exact redaction
# ---------------------------------------------------------------------------
def process_pdf(file_bytes: bytes, active_entities: list[str], detect_raw_fn, custom_patterns: Optional[list] = None) -> tuple[bytes, list]:
    """Redact native-text PDFs at the word level.

    The old implementation used `page.search_for(substring)`, which silently
    failed on text split across lines / ligatures / soft hyphens (PII stayed
    visible while the report claimed it was masked) and over-redacted any
    other occurrence *containing* the substring (flagging "John" also blacked
    out "Johnston"). Here we extract word boxes ourselves, rebuild the page
    text from those words, run detection on it, and redact exactly the boxes
    whose words intersect a detection span. No silent failures: every flagged
    span maps to concrete rectangles or is reported as missed.
    """
    report = []
    out_io = io.BytesIO()

    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            raw_words = page.get_text("words", sort=True)  # (x0,y0,x1,y1,word,block,line,word_no)
            if not raw_words:
                continue

            # Rebuild page text from the extracted words and remember where
            # each word lives in that string (offsets are exact by construction).
            parts = []
            word_spans = []  # (start, end, index)
            pos = 0
            for wi, w in enumerate(raw_words):
                if wi > 0:
                    parts.append(" ")
                    pos += 1
                start = pos
                parts.append(w[4])
                pos += len(w[4])
                word_spans.append((start, pos, wi))
            page_text = "".join(parts)
            if not page_text.strip():
                continue

            language = pii_engine.resolve_language(page_text)
            name_lines = _pdf_name_lines(raw_words, word_spans) if page.number == 0 else []
            results = detect_raw_fn(page_text, active_entities, custom_patterns, language=language, name_lines=name_lines)
            if not results:
                continue

            found_types = set()
            missed = []
            redacted_indices = set()
            for res in results:
                found_types.add(res.entity_type)
                overlapped = {wi for ws, we, wi in word_spans if ws < res.end and we > res.start}
                if overlapped:
                    redacted_indices |= overlapped
                else:
                    # Span maps to no word box (should not happen with word-built
                    # text) — report it honestly instead of claiming success.
                    missed.append((res.entity_type, page_text[res.start:res.end]))

            if redacted_indices:
                for wi in sorted(redacted_indices):
                    w = raw_words[wi]
                    rect = fitz.Rect(w[0] - 1, w[1] - 1, w[2] + 1, w[3] + 1)
                    page.add_redact_annot(rect, fill=(0, 0, 0))
                page.apply_redactions()

            if found_types:
                report.append({
                    "text": "PDF Page Content",
                    "pii_types": sorted(found_types),
                })

            for ent, snippet in missed:
                report.append({"text": f"PDF (unmapped span: {ent})", "pii_types": [ent], "snippet": snippet})

        doc.save(out_io)

    return out_io.getvalue(), report


# ---------------------------------------------------------------------------
# DOCX — run-level redaction (formatting preserved) + headers/footers
# ---------------------------------------------------------------------------
def _para_text(para) -> str:
    return "".join(run.text for run in para.runs)


def _apply_spans_to_paragraph(para, spans, masking_style):
    """Replace flagged spans inside a paragraph at run granularity.

    Editing runs directly (instead of `para.text = masked`) keeps every
    unaffected run's formatting (bold / italic / font / color) intact — the
    old implementation collapsed each paragraph to a single run and destroyed
    all formatting. Spans that cross run boundaries are handled by applying
    the replacement to the tail of the first run and clearing the covered
    parts of the following runs.
    """
    runs = para.runs
    if not runs:
        return
    run_starts = []
    pos = 0
    for r in runs:
        run_starts.append(pos)
        pos += len(r.text)
    para_len = pos

    # Collect per-run edits as (local_start, local_end, token), then apply
    # each run from the end backwards so earlier local offsets stay valid.
    edits = {i: [] for i in range(len(runs))}
    for s, e, ent in spans:
        s = max(0, s)
        e = min(para_len, e)
        if e <= s:
            continue
        token = _mask_token(ent, masking_style)
        # Runs overlapping [s, e)
        first_run = None
        last_run = None
        for i, rs in enumerate(run_starts):
            re_ = rs + len(runs[i].text)
            if rs < e and re_ > s:
                if first_run is None:
                    first_run = i
                last_run = i
        if first_run is None:
            continue
        if first_run == last_run:
            rs = run_starts[first_run]
            edits[first_run].append((s - rs, e - rs, token))
        else:
            # Token goes in the tail of the first run; clear the rest.
            rs = run_starts[first_run]
            edits[first_run].append((s - rs, len(runs[first_run].text), token))
            for i in range(first_run + 1, last_run):
                edits[i].append((0, len(runs[i].text), ""))
            lrs = run_starts[last_run]
            edits[last_run].append((0, e - lrs, ""))

    for i, run_edits in edits.items():
        if not run_edits:
            continue
        run = runs[i]
        text = run.text
        for ls, le, token in sorted(run_edits, key=lambda x: x[0], reverse=True):
            text = text[:ls] + token + text[le:]
        run.text = text


def _process_docx_paragraph(para, active_entities, detect_raw_fn, masking_style, custom_patterns, report, report_label, name_lines=None):
    text = _para_text(para)
    if not text.strip():
        return
    language = pii_engine.resolve_language(text)
    results = detect_raw_fn(text, active_entities, custom_patterns, language=language, name_lines=name_lines)
    if not results:
        return
    spans = [(r.start, r.end, r.entity_type) for r in results]
    _apply_spans_to_paragraph(para, spans, masking_style)
    report.append({"text": report_label, "pii_types": sorted({r.entity_type for r in results})})


def process_docx(file_bytes: bytes, active_entities: list[str], detect_raw_fn, masking_style: str = "LABEL", custom_patterns: Optional[list] = None) -> tuple[bytes, list]:
    doc = Document(io.BytesIO(file_bytes))
    report = []

    paragraphs = list(doc.paragraphs)
    # Header-block name rule: first paragraph looks like a name and the next
    # paragraphs carry contact info -> PERSON candidate for that paragraph.
    first_name_lines = None
    if paragraphs:
        if pii_engine.name_line_tokens_ok(paragraphs[0].text.split()):
            contact = " ".join(p.text for p in paragraphs[1:4])
            if _CONTACT_IN_LINE_RE.search(contact):
                first_name_lines = [(0, len(paragraphs[0].text))]

    for idx, para in enumerate(paragraphs):
        _process_docx_paragraph(para, active_entities, detect_raw_fn, masking_style, custom_patterns, report, "DOCX Paragraph",
                                name_lines=first_name_lines if idx == 0 else None)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_docx_paragraph(para, active_entities, detect_raw_fn, masking_style, custom_patterns, report, "DOCX Table Cell")

    # PII often lives in headers/footers (company name, contact email) —
    # the old implementation never scanned them, so it leaked.
    for section in doc.sections:
        for container in (section.header, section.footer):
            for para in container.paragraphs:
                _process_docx_paragraph(para, active_entities, detect_raw_fn, masking_style, custom_patterns, report, "DOCX Header/Footer")
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            _process_docx_paragraph(para, active_entities, detect_raw_fn, masking_style, custom_patterns, report, "DOCX Header/Footer")

    out_io = io.BytesIO()
    doc.save(out_io)
    return out_io.getvalue(), report


# ---------------------------------------------------------------------------
# Images — local OCR + span-exact box redaction
# ---------------------------------------------------------------------------
def _word_boxes_in_line(line_text: str, line_box) -> list:
    """Estimate per-word boxes inside an OCR text line.

    RapidOCR returns one box per text *line* (not per word). To redact exactly
    the flagged words instead of the whole line, we slice the line's bounding
    box horizontally in proportion to each word's character length. Left- and
    right-edge positions are interpolated along the quad's top/bottom edges so
    slightly rotated lines stay approximately correct.
    """
    words = [(m.start(), m.end(), m.group()) for m in re.finditer(r"\S+", line_text)]
    if not words or len(line_text) == 0:
        return []
    total = len(line_text)
    tl, tr, br, bl = line_box
    boxes = []
    for ws, we, word in words:
        f1 = ws / total
        f2 = we / total
        x0 = min(tl[0] + f1 * (tr[0] - tl[0]), bl[0] + f1 * (br[0] - bl[0]))
        x1 = max(tl[0] + f2 * (tr[0] - tl[0]), bl[0] + f2 * (br[0] - bl[0]))
        y0 = min(tl[1], bl[1])
        y1 = max(tr[1], br[1])
        boxes.append((x0, y0, x1, y1, ws, we, word))
    return boxes


def mask_pii_in_image(image_bytes: bytes, active_entities: list[str], detect_raw_fn, custom_patterns: Optional[list] = None):
    import cv2
    import numpy as np

    nparr = np.frombuffer(image_bytes, np.uint8)
    img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
    if img is None:
        raise Exception("Failed to decode image using OpenCV.")

    engine = _get_ocr_engine()
    ocr_result, _ = engine(img)
    if not ocr_result:
        return image_bytes, []

    lines = [{"box": [tuple(map(float, p)) for p in entry[0]], "text": entry[1]} for entry in ocr_result]

    # Rebuild the full text from OCR lines and keep each line's offset window,
    # so Presidio spans map back to the exact line and then to word boxes.
    full_text_parts = []
    line_windows = []  # (start, end)
    pos = 0
    for ln in lines:
        line_windows.append((pos, pos + len(ln["text"])))
        full_text_parts.append(ln["text"])
        pos += len(ln["text"]) + 1
    full_text = "\n".join(ln["text"] for ln in lines)

    language = pii_engine.resolve_language(full_text)
    name_lines = []
    if lines:
        if pii_engine.name_line_tokens_ok(lines[0]["text"].split()):
            contact = " ".join(ln["text"] for ln in lines[1:4])
            if _CONTACT_IN_LINE_RE.search(contact):
                s, e = line_windows[0]
                name_lines.append((s, e))
    results = detect_raw_fn(full_text, active_entities, custom_patterns, language=language, name_lines=name_lines)
    if not results:
        return image_bytes, []

    masked = img.copy()
    report = []
    redacted_regions = []
    for res in results:
        s, e = res.start, res.end
        for li, (ls, le) in enumerate(line_windows):
            part_s = max(s, ls) - ls
            part_e = min(e, le) - ls
            if part_e <= part_s:
                continue
            line = lines[li]
            for (x0, y0, x1, y1, ws, we, word) in _word_boxes_in_line(line["text"], line["box"]):
                if ws < part_e and we > part_s:
                    pad = 1
                    p0 = (int(x0 - pad), int(y0 - pad))
                    p1 = (int(x1 + pad), int(y1 + pad))
                    cv2.rectangle(masked, p0, p1, (0, 0, 0), -1)
                    redacted_regions.append((p0, p1))
                    report.append({
                        "text": word,
                        "pii_types": sorted({res.entity_type}),
                    })

    if not report:
        return image_bytes, []

    success, buffer = cv2.imencode(".jpg", masked, [cv2.IMWRITE_JPEG_QUALITY, 95])
    return buffer.tobytes(), report
